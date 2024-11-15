from docx import Document

# Create a new, detailed Word document
doc = Document()

# Title Page
doc.add_heading('Have Our Fundamental Rights Been Successful in Restoring the Idea of Equality and Liberty in India?', 0)
doc.add_paragraph('Submitted by: [Your Name]\nInstitution: [Your Institution]\nDate: [Date]')

# Acknowledgement
doc.add_page_break()
doc.add_heading('Acknowledgement', level=1)
doc.add_paragraph(
    "I express my heartfelt gratitude to my mentors and teachers for their guidance throughout this project. "
    "I am also thankful to my family and friends for their constant support. Lastly, I appreciate the insights "
    "provided by various authors, publications, and online platforms that greatly contributed to the completion of this work."
)

# Table of Contents
doc.add_page_break()
doc.add_heading('Table of Contents', level=1)
toc_content = [
    '1. Introduction .................................................... 4',
    '2. Fundamental Rights Overview ........................ 6',
    '3. Case Study: Right to Equality ............................ 8',
    '4. Case Study: Right to Freedom and Liberty ............. 11',
    '5. Role of Judiciary .............................................. 14',
    '6. Role of Civil Society and Media ......................... 16',
    '7. Limitations ....................................................... 17',
    '8. Success Indicators and Improvements .................. 18',
    '9. Conclusion ....................................................... 19',
    '10. Recommendations ......................................... 20',
    '11. Bibliography .................................................. 21'
]
for line in toc_content:
    doc.add_paragraph(line)

# Introduction
doc.add_page_break()
doc.add_heading('Introduction', level=1)
doc.add_paragraph(
    "The adoption of the Indian Constitution in 1950 was a watershed moment for the nation, ensuring justice, "
    "liberty, and equality in a country deeply rooted in social hierarchies and colonial legacy. Fundamental Rights, "
    "enshrined in Part III of the Constitution, are the pillars supporting democracy and human dignity. However, "
    "despite these guarantees, societal challenges persist, raising questions about the success of these rights in "
    "ensuring genuine equality and liberty.\n\n"
    "This project aims to critically evaluate the impact of Fundamental Rights on the lives of citizens. We will explore "
    "how these rights have been used to address inequality, promote liberty, and protect vulnerable groups. Additionally, "
    "the role of institutions such as the judiciary, media, and civil society in upholding these ideals will be analyzed."
)

# Fundamental Rights Overview
doc.add_page_break()
doc.add_heading('Fundamental Rights Overview', level=1)
doc.add_paragraph(
    "The Fundamental Rights provided by the Indian Constitution aim to safeguard the dignity and freedom of individuals. "
    "They form the foundation of Indian democracy and ensure that citizens can live without fear of discrimination, exploitation, "
    "or repression. The six core Fundamental Rights include: "
)

rights = [
    "Right to Equality (Articles 14-18): Ensures equal treatment before the law and prohibits discrimination.",
    "Right to Freedom (Articles 19-22): Guarantees personal freedoms, such as speech, assembly, and movement.",
    "Right against Exploitation (Articles 23-24): Prohibits human trafficking, forced labor, and child labor.",
    "Right to Freedom of Religion (Articles 25-28): Protects religious beliefs and practices.",
    "Cultural and Educational Rights (Articles 29-30): Safeguards minority cultures and educational institutions.",
    "Right to Constitutional Remedies (Article 32): Allows citizens to approach courts for rights violations."
]
for right in rights:
    doc.add_paragraph(f"- {right}", style='List Bullet')

# Expanded Case Study: Right to Equality
doc.add_page_break()
doc.add_heading('Case Study – Right to Equality in India', level=1)
doc.add_heading('1. Abolition of Untouchability', level=2)
doc.add_paragraph(
    "Article 17 of the Constitution abolished untouchability, a practice that had long stigmatized lower-caste groups, "
    "particularly Dalits. The government reinforced this provision through the SC/ST (Prevention of Atrocities) Act, "
    "criminalizing acts of discrimination against marginalized communities. However, incidents of caste-based violence, "
    "such as the Khairlanji massacre, highlight that discrimination still persists at both social and institutional levels."
)

doc.add_heading('2. Reservation Policies and Their Impact', level=2)
doc.add_paragraph(
    "The introduction of reservations in education, employment, and political institutions has significantly enhanced "
    "opportunities for Scheduled Castes (SCs), Scheduled Tribes (STs), and Other Backward Classes (OBCs). While these policies "
    "have helped uplift many individuals, debates around the 'creamy layer' and over-representation raise concerns about fairness. "
    "Furthermore, communities such as Dalit women face a double burden of caste and gender discrimination."
)

doc.add_heading('3. Gender Equality – Successes and Challenges', level=2)
doc.add_paragraph(
    "While India has made progress toward gender equality through measures like the Maternity Benefit Act and the Domestic "
    "Violence Act, significant gaps remain. Issues such as wage disparities, domestic violence, and low political participation "
    "reflect the ongoing struggle for women's empowerment. Landmark judgments like Vishaka vs. State of Rajasthan have shaped the "
    "legal landscape, but enforcement remains inconsistent."
)

# Expanded Case Study: Right to Freedom and Liberty
doc.add_page_break()
doc.add_heading('Case Study – Right to Freedom and Liberty', level=1)
doc.add_heading('1. Decriminalization of Section 377', level=2)
doc.add_paragraph(
    "The Supreme Court's 2018 verdict decriminalizing same-sex relationships was a historic moment for the LGBTQ+ community. "
    "By striking down Section 377, the court affirmed that personal liberty includes the right to love and live freely, irrespective "
    "of gender and sexual orientation. However, the struggle for equal marriage rights and social acceptance continues."
)

doc.add_heading('2. Freedom of Expression vs. State Censorship', level=2)
doc.add_paragraph(
    "Freedom of speech is a cornerstone of democracy, but it faces challenges in India. Laws like the sedition act have been used "
    "to suppress dissent, raising concerns about state overreach. The arrest of activists and journalists under sedition charges, "
    "alongside internet shutdowns during protests, illustrates the conflict between security concerns and individual freedoms."
)

# Role of Judiciary
doc.add_page_break()
doc.add_heading('Role of the Judiciary in Safeguarding Fundamental Rights', level=1)
doc.add_paragraph(
    "The judiciary plays a crucial role in interpreting and enforcing Fundamental Rights. Landmark cases, such as Kesavananda Bharati, "
    "have established that the basic structure of the Constitution cannot be altered, ensuring the continuity of these rights. "
    "Public Interest Litigations (PILs) have empowered citizens to approach the courts directly, but concerns about judicial delays "
    "and overreach persist."
)

# Limitations
doc.add_page_break()
doc.add_heading('Limitations', level=1)
doc.add_paragraph(
    "Despite the constitutional safeguards, several limitations hinder the realization of equality and liberty. Social and economic "
    "inequalities persist, and marginalized communities often struggle to access justice. Additionally, regional disparities, gender gaps, "
    "and censorship raise questions about the effectiveness of these rights."
)

# Recommendations
doc.add_page_break()
doc.add_heading('Recommendations', level=1)
doc.add_paragraph(
    "1. Strengthen awareness campaigns to educate citizens about their rights.\n"
    "2. Reform outdated laws, such as the sedition act, to prevent misuse.\n"
    "3. Enhance the capacity of the judiciary to address delays and improve access to justice.\n"
    "4. Promote inclusive policies to ensure equal opportunities across all regions and communities."
)

# Conclusion
doc.add_page_break()
doc.add_heading('Conclusion', level=1)
doc.add_paragraph(
    "While the Indian Constitution provides a robust framework for equality and liberty, the journey toward realizing these ideals is ongoing. "
    "Continuous efforts are required from the government, judiciary, civil society, and citizens to address the existing challenges. Only through "
    "collective action can India ensure that the promise of Fundamental Rights becomes a lived reality for all."
)

# Bibliography
doc.add_page_break()
doc.add_heading('Bibliography', level=1)
doc.add_paragraph(
    "1. Austin, Granville. 'The Indian Constitution: Cornerstone of a Nation'.\n"
    "2. Supreme Court Judgments: Kesavananda Bharati, Navtej Singh Johar, and Maneka Gandhi.\n"
    "3. Reports from the National Human Rights Commission (NHRC).\n"
    "4. News Articles: The Hindu, Indian Express, Times of India."
)

# Save the final document
output_path = "F:/PYTHON/Neural Network/Comprehensive_Fundamental_Rights_Project.docx"
doc.save(output_path)

# Provide the file path for download
output_path
